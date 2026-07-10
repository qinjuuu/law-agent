"""
Word 文档生成工具模块
生成消费维权投诉信和格式条款审查报告
"""
import os
import sys
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain.tools import tool

from config import WORD_REPORTS_DIR

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")


def _safe_word_path(filename: str) -> str:
    """确保 Word 文件保存在 WORD_REPORTS_DIR 下，自动补全 .docx 后缀"""
    if not filename.endswith(".docx"):
        filename += ".docx"
    abs_path = os.path.normpath(os.path.join(WORD_REPORTS_DIR, filename))
    if not abs_path.startswith(os.path.normpath(WORD_REPORTS_DIR)):
        raise ValueError(f"路径越界: {filename}")
    return abs_path


def _setup_doc_style(doc: Document):
    """设置文档默认样式（中文字体、行距）"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(12)
    # 设置中文字体
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(6)


@tool
def create_complaint_report(
    title: str,
    complainant: str,
    respondent: str,
    complaint_content: List[str],
    legal_basis: List[str],
    demands: List[str],
    filename: Optional[str] = None,
    confirm: bool = False,
) -> str:
    """
    生成一份消费维权投诉信 Word 文档

    参数:
        title: 投诉信标题
        complainant: 投诉人信息（姓名、联系方式）
        respondent: 被投诉方信息（商家名称、地址）
        complaint_content: 投诉事实描述，每个元素为一个段落
        legal_basis: 法律依据列表，引用的相关法条
        demands: 诉求列表，如退款、赔偿、道歉等
        filename: 输出文件名，未提供时自动用时间戳生成
        confirm: 是否确认操作

    返回:
        操作结果信息

    适用场景:
        用户需要生成正式的消费者投诉信时调用
    """
    try:
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"complaint_{timestamp}.docx"

        file_path = _safe_word_path(filename)

        if os.path.exists(file_path) and not confirm:
            return f"文件 {filename} 已存在，如需覆盖，请设置 confirm 为 True"
        if not confirm:
            return f"请确认创建文件 {filename}，如确认需将 confirm 设置为 True"

        doc = Document()
        _setup_doc_style(doc)

        # 标题
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 投诉人信息
        doc.add_heading("投诉人信息", level=2)
        doc.add_paragraph(complainant)

        # 被投诉方信息
        doc.add_heading("被投诉方信息", level=2)
        doc.add_paragraph(respondent)

        # 投诉事实
        doc.add_heading("投诉事实", level=2)
        for para in complaint_content:
            doc.add_paragraph(para)

        # 法律依据
        doc.add_heading("法律依据", level=2)
        for law in legal_basis:
            p = doc.add_paragraph(law)
            p.paragraph_format.left_indent = Cm(0.74)

        # 诉求
        doc.add_heading("诉求", level=2)
        for i, demand in enumerate(demands, 1):
            doc.add_paragraph(f"{i}. {demand}")

        # 结尾
        doc.add_paragraph("")
        doc.add_paragraph("此致")
        doc.add_paragraph("相关市场监督管理部门")
        doc.add_paragraph("")
        p = doc.add_paragraph(f"投诉人（签字）: ___________")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p = doc.add_paragraph(f"日期: {datetime.now().strftime('%Y年%m月%d日')}")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(file_path)
        return f"投诉信已生成: {filename}，共 {len(complaint_content) + len(legal_basis) + len(demands) + 6} 个段落"
    except ValueError:
        return f"安全问题: {filename}"
    except Exception as e:
        return f"生成失败: {str(e)}"


@tool
def create_review_report(
    contract_title: str,
    review_results: List[str],
    risk_items: List[str],
    suggestions: List[str],
    risk_level: str,
    filename: Optional[str] = None,
    confirm: bool = False,
) -> str:
    """
    生成一份格式条款审查报告 Word 文档

    参数:
        contract_title: 被审查合同/条款名称
        review_results: 审查结论列表，每个元素为一个分析段落
        risk_items: 风险条款列表，标注具体风险点
        suggestions: 修改建议列表
        risk_level: 整体风险等级（低/中/高）
        filename: 输出文件名，未提供时自动用时间戳生成
        confirm: 是否确认操作

    返回:
        操作结果信息

    适用场景:
        用户需要生成消费合同/格式条款审查报告时调用
    """
    try:
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            filename = f"review_{timestamp}.docx"

        file_path = _safe_word_path(filename)

        if os.path.exists(file_path) and not confirm:
            return f"文件 {filename} 已存在，如需覆盖，请设置 confirm 为 True"
        if not confirm:
            return f"请确认创建文件 {filename}，如确认需将 confirm 设置为 True"

        doc = Document()
        _setup_doc_style(doc)

        # 标题
        heading = doc.add_heading(f"格式条款审查报告", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"审查对象: {contract_title}")
        doc.add_paragraph(f"审查日期: {datetime.now().strftime('%Y年%m月%d日')}")
        doc.add_paragraph(f"整体风险等级: {risk_level}")
        doc.add_paragraph("")

        # 风险等级颜色标记
        risk_colors = {"低": RGBColor(0x00, 0x80, 0x00), "中": RGBColor(0xFF, 0x80, 0x00), "高": RGBColor(0xFF, 0x00, 0x00)}
        risk_color = risk_colors.get(risk_level, RGBColor(0x00, 0x00, 0x00))

        # 审查结论
        doc.add_heading("一、审查结论", level=2)
        for para in review_results:
            doc.add_paragraph(para)

        # 风险条款
        doc.add_heading("二、风险条款识别", level=2)
        if risk_items:
            for i, item in enumerate(risk_items, 1):
                p = doc.add_paragraph(f"风险{i}: {item}")
                p.runs[0].font.color.rgb = risk_color
        else:
            doc.add_paragraph("未发现明显风险条款")

        # 修改建议
        doc.add_heading("三、修改建议", level=2)
        for i, suggestion in enumerate(suggestions, 1):
            doc.add_paragraph(f"{i}. {suggestion}")

        # 结尾
        doc.add_paragraph("")
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.save(file_path)
        return f"审查报告已生成: {filename}，风险等级 {risk_level}，共 {len(review_results) + len(risk_items) + len(suggestions) + 4} 个段落"
    except ValueError:
        return f"安全问题: {filename}"
    except Exception as e:
        return f"生成失败: {str(e)}"
